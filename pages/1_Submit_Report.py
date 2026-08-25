import streamlit as st
import sqlite3
import json
import io
import pandas as pd
from google import genai
from datetime import datetime, date
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Submit Report", page_icon="📝", layout="centered")

DB_PATH = Path("project_data.db")

# ── PASSWORD CHECK ────────────────────────────────────────────────────────
def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if st.session_state.authenticated:
        return True
    st.title("Team Login")
    st.markdown("Only team members can submit reports.")
    st.divider()
    password = st.text_input("Enter team password", type="password")
    if st.button("Login", use_container_width=True, type="primary"):
        if password == st.secrets["DASHBOARD_PASSWORD"]:
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Incorrect password. Please try again.")
    return False

if not check_password():
    st.stop()

# ── DB ────────────────────────────────────────────────────────────────────
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS reports (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            submitted_at TEXT,
            member       TEXT,
            week_label   TEXT,
            raw_report   TEXT,
            tasks_json   TEXT
        )
    """)
    conn.commit()
    return conn

# ── READ TEXT FROM UPLOADED DOCX ─────────────────────────────────────────
def extract_text_from_docx(uploaded_file) -> str:
    doc = Document(uploaded_file)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def parse_report_with_ai(member, week, report_text):
    client = genai.Client(api_key=st.secrets["GEMINI_API_KEY"])
    prompt = f"""
You are extracting structured work data from a student weekly project report.
Student: {member}
Week: {week}
Report: \"\"\"{report_text}\"\"\"

Extract every distinct task. Return a JSON array:
[{{"task": "short name max 10 words", "category": "one of Coding/Research/Engineering/Debugging/Writing/Meeting/Other", "hours": number, "evidence": "link or empty string"}}]
Return ONLY the JSON array. No explanation. No markdown. No backticks.
""".strip()

    candidate_models = ["gemini-2.5-flash", "gemini-1.5-flash"]
    last_error = None

    for model_name in candidate_models:
        try:
            response = client.models.generate_content(model=model_name, contents=prompt)
            raw = response.text.strip()
            if raw.startswith("```json"):
                raw = raw[7:]
            elif raw.startswith("```"):
                raw = raw[3:]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()
            return json.loads(raw)
        except Exception as e:
            last_error = e
            continue

    if last_error:
        raise last_error

def save_report(member, week_label, raw_report, tasks):
    conn = get_db()
    conn.execute(
        "INSERT INTO reports (submitted_at, member, week_label, raw_report, tasks_json) VALUES (?,?,?,?,?)",
        (datetime.now().isoformat(), member, week_label, raw_report, json.dumps(tasks)),
    )
    conn.commit()
    conn.close()

# ── WEEK COMPARISON GRAPHS ────────────────────────────────────────────────
def show_week_comparison(member):
    conn = get_db()
    rows = conn.execute(
        "SELECT tasks_json, week_label FROM reports WHERE member=? ORDER BY submitted_at ASC",
        (member,)
    ).fetchall()
    conn.close()

    if len(rows) < 2:
        st.info("Submit at least 2 weekly reports to see week-on-week comparison graphs.")
        return

    weekly_hours  = []
    weekly_tasks  = []
    weekly_labels = []
    weekly_cats   = []

    for row in rows:
        tasks = json.loads(row[0])
        label = row[1].replace("Week ending ", "")
        weekly_labels.append(label)
        weekly_hours.append(sum(float(t.get("hours", 0)) for t in tasks))
        weekly_tasks.append(len(tasks))
        for t in tasks:
            weekly_cats.append({
                "Week": label,
                "Category": t.get("category", "Other"),
                "Hours": float(t.get("hours", 0))
            })

    curr_hours = weekly_hours[-1]
    prev_hours = weekly_hours[-2]
    curr_count = weekly_tasks[-1]
    prev_count = weekly_tasks[-2]
    hour_diff  = curr_hours - prev_hours
    task_diff  = curr_count - prev_count

    st.subheader("Weekly Progress Comparison")
    st.divider()

    c1, c2, c3 = st.columns(3)
    c1.metric("Hours This Week", f"{curr_hours:.1f} hrs", delta=f"{hour_diff:+.1f} hrs vs last week")
    c2.metric("Tasks This Week", curr_count,              delta=f"{task_diff:+d} tasks vs last week")
    c3.metric("Weeks Tracked",   len(rows))

    if hour_diff > 0:
        st.success(f"You put in **{hour_diff:.1f} more hours** than last week. Great progress!")
    elif hour_diff < 0:
        st.warning(f"You put in **{abs(hour_diff):.1f} fewer hours** than last week.")
    else:
        st.info("Same number of hours as last week.")

    st.write("")

    st.markdown("**Hours logged per week**")
    hours_df = pd.DataFrame({"Week": weekly_labels, "Hours": weekly_hours}).set_index("Week")
    st.bar_chart(hours_df, use_container_width=True)

    st.markdown("**Tasks completed per week**")
    tasks_df = pd.DataFrame({"Week": weekly_labels, "Tasks": weekly_tasks}).set_index("Week")
    st.bar_chart(tasks_df, use_container_width=True)

    if weekly_cats:
        st.markdown("**Effort by category across all weeks**")
        cats_df = pd.DataFrame(weekly_cats)
        pivot   = cats_df.groupby(["Week", "Category"])["Hours"].sum().unstack(fill_value=0)
        st.bar_chart(pivot, use_container_width=True)

# ── WORD DOC GENERATOR ────────────────────────────────────────────────────
def generate_word_report(member, week_label, raw_report, tasks):
    doc = Document()

    title = doc.add_heading("Final Year Project - Weekly Progress Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{member}  |  {week_label}  |  {datetime.now().strftime('%d %b %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    doc.add_heading("Summary", level=1)
    total_hours = sum(float(t.get("hours", 0)) for t in tasks)
    categories  = list(set(t.get("category") for t in tasks))

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = f"Total Hours\n{total_hours:.1f} hrs"
    hdr[1].text = f"Tasks Completed\n{len(tasks)}"
    hdr[2].text = f"Categories\n{', '.join(categories)}"
    doc.add_paragraph()

    doc.add_heading("Weekly Report (Student's Own Words)", level=1)
    doc.add_paragraph(raw_report)
    doc.add_paragraph()

    doc.add_heading("Extracted Task Breakdown", level=1)
    task_tbl = doc.add_table(rows=1, cols=4)
    task_tbl.style = "Table Grid"
    hdrs = task_tbl.rows[0].cells
    hdrs[0].text = "Task"
    hdrs[1].text = "Category"
    hdrs[2].text = "Hours"
    hdrs[3].text = "Evidence"

    for t in tasks:
        row = task_tbl.add_row().cells
        row[0].text = t.get("task", "")
        row[1].text = t.get("category", "")
        row[2].text = str(t.get("hours", ""))
        row[3].text = t.get("evidence", "") or "-"

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Generated by Final Year Project Dashboard")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ── SHARED SUBMISSION HANDLER ─────────────────────────────────────────────
def handle_submission(member, week_label, week, report_text):
    with st.spinner("AI is reading your report and extracting your work..."):
        try:
            tasks = parse_report_with_ai(member, week_label, report_text)
            save_report(member, week_label, report_text, tasks)
        except json.JSONDecodeError:
            st.error("AI returned unexpected format. Please try again.")
            st.stop()
        except Exception as e:
            st.error(f"Something went wrong: {e}")
            st.stop()

    st.success(f"Report submitted! {len(tasks)} tasks extracted.")
    st.divider()

    st.subheader("Tasks AI Extracted")
    for t in tasks:
        c1, c2, c3 = st.columns([4, 2, 1])
        c1.markdown(f"**{t.get('task', '-')}**")
        c2.markdown(f"`{t.get('category', '-')}`")
        c3.markdown(f"**{t.get('hours', 0)} hrs**")

    total = sum(float(t.get("hours", 0)) for t in tasks)
    st.metric("Total hours logged this week", f"{total:.1f} hrs")
    st.divider()

    show_week_comparison(member)
    st.divider()

    st.subheader("Download This Report")
    word_buf = generate_word_report(member, week_label, report_text, tasks)
    filename = f"Weekly_Report_{member}_{week.strftime('%d_%b_%Y')}.docx".replace(" ", "_")
    st.download_button(
        label="Download as Word Document (.docx)",
        data=word_buf,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )
    st.info("Share this Word file with your guide or keep it for your records.")

# ── UI ────────────────────────────────────────────────────────────────────
col1, col2 = st.columns([5, 1])
with col1:
    st.title("Submit Weekly Report")
with col2:
    st.write("")
    if st.button("Logout"):
        st.session_state.authenticated = False
        st.rerun()

st.markdown("Submit your weekly report by typing it or uploading a Word file. AI will extract your tasks automatically.")
st.divider()

member     = st.selectbox("Your Name", options=["Shubham", "Ankit Kumar"])
week       = st.date_input("Week Ending On", value=date.today())
week_label = f"Week ending {week.strftime('%d %b %Y')}"

# ── TWO TABS ──────────────────────────────────────────────────────────────
tab1, tab2 = st.tabs(["✏️  Type your report", "📄  Upload Word file (.docx)"])

with tab1:
    report_text = st.text_area(
        "Your weekly report",
        height=250,
        placeholder=(
            "Write what you did this week. For example:\n\n"
            "This week I spent around 4 hours reading papers on transformer architectures "
            "and took notes on the attention mechanism. I also spent 3 hours debugging the "
            "data pipeline — the issue was a mismatch in tensor shapes during batching, "
            "which I fixed by reshaping the input. On Friday I trained the baseline CNN "
            "model for 2 hours and got 78% accuracy on the validation set."
        ),
        key="typed_report"
    )
    if st.button("Submit Report", use_container_width=True, type="primary", key="submit_text"):
        if not report_text.strip():
            st.error("Please write something before submitting.")
        else:
            handle_submission(member, week_label, week, report_text)

with tab2:
    st.markdown("Write your weekly report in a Word document and upload it here.")
    st.caption("Supported format: .docx only")

    uploaded_file = st.file_uploader(
        "Upload your Word report",
        type=["docx"],
        key="uploaded_docx"
    )

    if uploaded_file is not None:
        with st.spinner("Reading your Word document..."):
            try:
                extracted_text = extract_text_from_docx(uploaded_file)
            except Exception as e:
                st.error(f"Could not read the file: {e}")
                st.stop()

        st.success(f"Word file read successfully — {len(extracted_text.split())} words found.")

        with st.expander("Preview extracted text"):
            st.text(extracted_text[:1000] + ("..." if len(extracted_text) > 1000 else ""))

        if st.button("Submit This Report", use_container_width=True, type="primary", key="submit_docx"):
            if not extracted_text.strip():
                st.error("The Word file appears to be empty. Please check and re-upload.")
            else:
                handle_submission(member, week_label, week, extracted_text)
